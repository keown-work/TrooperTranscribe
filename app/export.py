# Copyright (c) 2026 Casey Keown. All rights reserved. Proprietary and confidential.

"""
Export transcript segments to TXT, DOCX, PDF, SRT, VTT, and CSV.
"""

import io
from typing import Tuple


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def export_transcript(
    segments: list, fmt: str, filename: str = "transcript"
) -> Tuple[bytes, str, str]:
    """
    Convert segments to the requested format.
    Returns (raw_bytes, mime_type, file_extension).
    """
    handlers = {
        "txt":  _export_txt,
        "docx": _export_docx,
        "pdf":  _export_pdf,
        "srt":  _export_srt,
        "vtt":  _export_vtt,
        "csv":  _export_csv,
    }
    if fmt not in handlers:
        raise ValueError(f"Unsupported format: {fmt}")
    return handlers[fmt](segments, filename)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _ts_srt(s: float) -> str:
    h, rem = divmod(int(s), 3600)
    m, sec = divmod(rem, 60)
    ms = int(round((s % 1) * 1000))
    return f"{h:02d}:{m:02d}:{sec:02d},{ms:03d}"


def _ts_vtt(s: float) -> str:
    return _ts_srt(s).replace(",", ".")


def _ts_readable(s: float) -> str:
    h, rem = divmod(int(s), 3600)
    m, sec = divmod(rem, 60)
    if h:
        return f"{h:02d}:{m:02d}:{sec:02d}"
    return f"{m:02d}:{sec:02d}"


# ---------------------------------------------------------------------------
# Format implementations
# ---------------------------------------------------------------------------

def _export_txt(segments: list, filename: str) -> Tuple[bytes, str, str]:
    lines = []
    current_speaker = None
    for seg in segments:
        spk = seg.get("speaker", "SPEAKER_01")
        if spk != current_speaker:
            if lines:
                lines.append("")
            lines.append(f"[{spk}]")
            current_speaker = spk
        ts = _ts_readable(seg["start"])
        lines.append(f"[{ts}]  {seg['text']}")
    return "\n".join(lines).encode("utf-8"), "text/plain; charset=utf-8", "txt"


def _export_docx(segments: list, filename: str) -> Tuple[bytes, str, str]:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Header
    title = doc.add_heading("TRANSCRIPTION REPORT", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(f"Source file: {filename}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    current_speaker = None
    for seg in segments:
        spk = seg.get("speaker", "SPEAKER_01")
        ts = f"[{_ts_readable(seg['start'])} — {_ts_readable(seg['end'])}]"

        if spk != current_speaker:
            sp = doc.add_paragraph()
            r = sp.add_run(spk)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            current_speaker = spk

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        r_ts = p.add_run(ts + "  ")
        r_ts.font.size = Pt(8)
        r_ts.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        r_text = p.add_run(seg["text"])
        r_text.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return (
        buf.read(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    )


def _export_pdf(segments: list, filename: str) -> Tuple[bytes, str, str]:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "KSPTitle", parent=styles["Heading1"], fontSize=14, spaceAfter=4
    )
    meta_style = ParagraphStyle(
        "KSPMeta", parent=styles["Normal"], fontSize=9, spaceAfter=12,
        textColor=colors.HexColor("#555555")
    )
    speaker_style = ParagraphStyle(
        "KSPSpeaker", parent=styles["Normal"], fontSize=9,
        fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=2,
        textColor=colors.HexColor("#222222")
    )
    ts_style = ParagraphStyle(
        "KSPTimestamp", parent=styles["Normal"], fontSize=7,
        textColor=colors.HexColor("#888888"), leftIndent=16, spaceAfter=1
    )
    text_style = ParagraphStyle(
        "KSPText", parent=styles["Normal"], fontSize=10,
        leftIndent=16, spaceAfter=3
    )

    story = [
        Paragraph("TRANSCRIPTION REPORT", title_style),
        Paragraph(f"Source: {filename}", meta_style),
    ]

    current_speaker = None
    for seg in segments:
        spk = seg.get("speaker", "SPEAKER_01")
        if spk != current_speaker:
            story.append(Paragraph(spk, speaker_style))
            current_speaker = spk
        story.append(
            Paragraph(
                f"[{_ts_readable(seg['start'])} \u2014 {_ts_readable(seg['end'])}]",
                ts_style,
            )
        )
        story.append(Paragraph(seg["text"], text_style))

    doc.build(story)
    buf.seek(0)
    return buf.read(), "application/pdf", "pdf"


def _export_srt(segments: list, filename: str) -> Tuple[bytes, str, str]:
    lines = []
    for i, seg in enumerate(segments, 1):
        spk = seg.get("speaker", "SPEAKER_01")
        lines += [
            str(i),
            f"{_ts_srt(seg['start'])} --> {_ts_srt(seg['end'])}",
            f"[{spk}] {seg['text']}",
            "",
        ]
    return "\n".join(lines).encode("utf-8"), "text/srt", "srt"


def _export_vtt(segments: list, filename: str) -> Tuple[bytes, str, str]:
    lines = ["WEBVTT", ""]
    for seg in segments:
        spk = seg.get("speaker", "SPEAKER_01")
        lines += [
            f"{_ts_vtt(seg['start'])} --> {_ts_vtt(seg['end'])}",
            f"<v {spk}>{seg['text']}",
            "",
        ]
    return "\n".join(lines).encode("utf-8"), "text/vtt", "vtt"


def _export_csv(segments: list, filename: str) -> Tuple[bytes, str, str]:
    import csv

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["speaker", "start_seconds", "end_seconds", "start_time", "end_time", "text"])
    for seg in segments:
        writer.writerow([
            seg.get("speaker", "SPEAKER_01"),
            seg.get("start", 0),
            seg.get("end", 0),
            _ts_readable(seg.get("start", 0)),
            _ts_readable(seg.get("end", 0)),
            seg.get("text", ""),
        ])
    return buf.getvalue().encode("utf-8"), "text/csv; charset=utf-8", "csv"
